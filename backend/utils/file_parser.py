# backend/utils/file_parser.py

import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageFilter, ImageEnhance
import pytesseract

pytesseract.pytesseract.tesseract_cmd = (
    r"C:\Users\Pragati\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
)


def parse_pdf(file_path: str) -> str:
    """
    Extract text from PDF using PyMuPDF.

    Uses block-level extraction sorted by vertical position
    to preserve reading order on multi-column resumes.
    """

    text = []

    pdf_document = fitz.open(file_path)

    try:
        for page in pdf_document:

            # get_text("blocks") returns:
            # (x0, y0, x1, y1, text, block_no, block_type)

            blocks = page.get_text("blocks")

            # sort by vertical position first,
            # then horizontal position

            blocks.sort(key=lambda b: (b[1], b[0]))

            for block in blocks:

                block_text = block[4].strip()

                if block_text:
                    text.append(block_text)

    finally:
        pdf_document.close()

    return "\n".join(text).strip()


def parse_docx(file_path: str) -> str:
    """
    Extract text from DOCX.

    Includes paragraphs and tables.
    """

    document = Document(file_path)

    text = []

    for paragraph in document.paragraphs:

        if paragraph.text.strip():
            text.append(paragraph.text)

    for table in document.tables:

        for row in table.rows:

            row_text = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if row_text:
                text.append(
                    " | ".join(row_text)
                )

    return "\n".join(text).strip()


def parse_image(file_path: str) -> str:
    """
    OCR image resume using Tesseract.

    Preprocesses the image before OCR:

    - Convert to grayscale
    - Upscale if small
    - Sharpen image
    - Enhance contrast
    """

    image = Image.open(file_path).convert("L")

    # upscale if image is smaller than 1800px wide

    if image.width < 1800:

        scale_factor = 1800 / image.width

        new_size = (
            int(image.width * scale_factor),
            int(image.height * scale_factor)
        )

        image = image.resize(
            new_size,
            Image.LANCZOS
        )

    # sharpen image

    image = image.filter(
        ImageFilter.SHARPEN
    )

    # increase contrast

    enhancer = ImageEnhance.Contrast(
        image
    )

    image = enhancer.enhance(2.0)

    # OCR configuration

    custom_config = "--oem 3 --psm 6"

    extracted_text = pytesseract.image_to_string(
        image,
        config=custom_config
    )

    result = extracted_text.strip()

    if not result:
        raise ValueError(
            "No text could be extracted from image."
        )

    return result